#!/usr/bin/env python3
"""
Assignment Project File Generator
----------------------------------
Capture a screen region -> LLM Vision writes a one-line step description
-> Screenshot + step auto-appended to a Word (.docx) file.

Supported platforms: Anthropic (Claude), OpenAI (GPT-4o), Google (Gemini)
"""

import os, sys, io, base64, json, datetime
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
from PIL import Image, ImageGrab
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
#  LLM Platform Registry
# ─────────────────────────────────────────────────────────────

PLATFORMS = {
    "Anthropic  (Claude)": {
        "key_label": "Anthropic API Key",
        "key_hint":  "Starts with:  sk-ant-...",
        "models": [
            "claude-opus-4-5",
            "claude-sonnet-4-5",
            "claude-haiku-3-5-20251001",
        ],
    },
    "OpenAI  (GPT-4o)": {
        "key_label": "OpenAI API Key",
        "key_hint":  "Starts with:  sk-...",
        "models": [
            "gpt-4o",
            "gpt-4o-mini",
            "gpt-4-turbo",
        ],
    },
    "Google  (Gemini)": {
        "key_label": "Google AI API Key",
        "key_hint":  "Starts with:  AIza...",
        "models": [
            "gemini-2.0-flash",
            "gemini-1.5-pro",
            "gemini-1.5-flash",
        ],
    },
}

STEP_PROMPT = (
    "Describe the action shown in this screenshot as a single step instruction. "
    "Rules: ONE sentence only. Start with a verb. Maximum 12 words. "
    "No step number prefix. No period at the end. Just the action itself."
)


# ─────────────────────────────────────────────────────────────
#  LLM call wrappers
# ─────────────────────────────────────────────────────────────

def call_anthropic(api_key, model, image_b64):
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=model,
        max_tokens=80,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/png",
                        "data": image_b64,
                    },
                },
                {"type": "text", "text": STEP_PROMPT},
            ],
        }],
    )
    return resp.content[0].text.strip().rstrip(".")


def call_openai(api_key, model, image_b64):
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        max_tokens=80,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{image_b64}",
                        "detail": "low",
                    },
                },
                {"type": "text", "text": STEP_PROMPT},
            ],
        }],
    )
    return resp.choices[0].message.content.strip().rstrip(".")


def call_gemini(api_key, model, image_bytes):
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    m = genai.GenerativeModel(model)
    pil_img = Image.open(io.BytesIO(image_bytes))
    resp = m.generate_content([STEP_PROMPT, pil_img])
    return resp.text.strip().rstrip(".")


def describe_screenshot(platform, api_key, model, image: Image.Image) -> str:
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    raw_bytes = buf.getvalue()
    b64 = base64.standard_b64encode(raw_bytes).decode()

    if "Anthropic" in platform:
        return call_anthropic(api_key, model, b64)
    elif "OpenAI" in platform:
        return call_openai(api_key, model, b64)
    elif "Google" in platform:
        return call_gemini(api_key, model, raw_bytes)
    else:
        raise ValueError(f"Unknown platform: {platform}")


# ─────────────────────────────────────────────────────────────
#  Screen-region selector overlay
# ─────────────────────────────────────────────────────────────

class RegionSelector:
    def __init__(self, root_window, callback):
        self.callback = callback
        self.start_x = self.start_y = 0
        self.rect_id = None

        self.full_screenshot = ImageGrab.grab(all_screens=True)

        self.win = tk.Toplevel(root_window)
        self.win.attributes("-fullscreen", True)
        self.win.attributes("-alpha", 0.35)
        self.win.attributes("-topmost", True)
        self.win.configure(bg="#000010")
        self.win.overrideredirect(True)

        self.canvas = tk.Canvas(
            self.win, cursor="crosshair", bg="#000010", highlightthickness=0
        )
        self.canvas.pack(fill="both", expand=True)

        sw = self.win.winfo_screenwidth()
        self.canvas.create_text(
            sw // 2, 45,
            text="Click and drag to select a region  |  ESC to cancel",
            fill="#00ff99",
            font=("Arial", 17, "bold"),
        )

        self.canvas.bind("<ButtonPress-1>", self._press)
        self.canvas.bind("<B1-Motion>", self._drag)
        self.canvas.bind("<ButtonRelease-1>", self._release)
        self.win.bind("<Escape>", lambda _: self._cancel())

    def _press(self, e):
        self.start_x, self.start_y = e.x, e.y

    def _drag(self, e):
        if self.rect_id:
            self.canvas.delete(self.rect_id)
        self.rect_id = self.canvas.create_rectangle(
            self.start_x, self.start_y, e.x, e.y,
            outline="#00ff99", width=2, fill="",
        )

    def _release(self, e):
        x1 = min(self.start_x, e.x)
        y1 = min(self.start_y, e.y)
        x2 = max(self.start_x, e.x)
        y2 = max(self.start_y, e.y)
        self.win.destroy()
        if x2 - x1 < 15 or y2 - y1 < 15:
            messagebox.showwarning("Too small", "Selection too small, please try again.")
            self.callback(None)
            return
        self.callback(self.full_screenshot.crop((x1, y1, x2, y2)))

    def _cancel(self):
        self.win.destroy()
        self.callback(None)


# ─────────────────────────────────────────────────────────────
#  Word document helpers
# ─────────────────────────────────────────────────────────────

def create_new_document(project_name, subject):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(project_name)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)

    if subject:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub_p.add_run(f"Subject: {subject}")
        sr.italic = True
        sr.font.size = Pt(12)
        sr.font.color.rgb = RGBColor(0x55, 0x55, 0x88)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(datetime.date.today().strftime("Date: %B %d, %Y"))
    dr.italic = True
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor(0x55, 0x55, 0x88)

    rule_p = doc.add_paragraph()
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C0392B")
    pBdr.append(bot)
    pPr.append(pBdr)

    doc.add_paragraph()
    return doc


def append_step(doc, step_num, description, image: Image.Image):
    step_p = doc.add_paragraph()
    step_r = step_p.add_run(f"Step {step_num}:  {description}")
    step_r.bold = True
    step_r.font.size = Pt(12)
    step_r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    max_px = 950
    w, h = image.size
    if w > max_px:
        image = image.resize((max_px, int(h * max_px / w)), Image.LANCZOS)

    buf = io.BytesIO()
    image.save(buf, format="PNG")
    buf.seek(0)

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(buf, width=Inches(5.5))

    div_p = doc.add_paragraph()
    pPr2 = div_p._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot2 = OxmlElement("w:bottom")
    bot2.set(qn("w:val"), "single")
    bot2.set(qn("w:sz"), "4")
    bot2.set(qn("w:space"), "1")
    bot2.set(qn("w:color"), "CCCCCC")
    pBdr2.append(bot2)
    pPr2.append(pBdr2)

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────
#  UI constants
# ─────────────────────────────────────────────────────────────

BG_DARK  = "#1a1a2e"
BG_CARD  = "#16213e"
ACCENT   = "#c0392b"
ACCENT2  = "#0f3460"
TXT_MAIN = "#eaeaea"
TXT_SUB  = "#8888aa"
TXT_OK   = "#27ae60"

SETTINGS_FILE = os.path.expanduser("~/.asgn_gen_v2.json")


def styled_button(parent, text, cmd, bg=ACCENT2, fg=TXT_MAIN, **kw):
    b = tk.Button(
        parent, text=text, command=cmd,
        bg=bg, fg=fg,
        activebackground=ACCENT, activeforeground="white",
        relief="flat", bd=0, cursor="hand2",
        font=("Arial", 10, "bold"), pady=9, **kw,
    )
    orig_bg = bg
    b.bind("<Enter>", lambda _: b.config(bg=ACCENT))
    b.bind("<Leave>", lambda _: b.config(bg=orig_bg))
    return b


def small_label(parent, text):
    return tk.Label(
        parent, text=text,
        bg=BG_DARK, fg=TXT_SUB,
        font=("Arial", 9), anchor="w",
    )


def text_entry(parent, var=None, show=None):
    kw = dict(bg=BG_CARD, fg=TXT_MAIN, insertbackground=TXT_MAIN,
              relief="flat", bd=8, font=("Arial", 10))
    if var:
        kw["textvariable"] = var
    if show:
        kw["show"] = show
    return tk.Entry(parent, **kw)


# ─────────────────────────────────────────────────────────────
#  Main Application
# ─────────────────────────────────────────────────────────────

class App:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Assignment Project Generator")
        self.root.geometry("460x670")
        self.root.resizable(False, False)
        self.root.configure(bg=BG_DARK)
        self.root.protocol("WM_DELETE_WINDOW", self._quit)

        self.platform_var = tk.StringVar(value=list(PLATFORMS.keys())[0])
        self.api_key_var  = tk.StringVar()
        self.model_var    = tk.StringVar()
        self.project_var  = tk.StringVar()
        self.subject_var  = tk.StringVar()
        self.status_var   = tk.StringVar(value="Configure platform and project, then click Start.")
        self.doc_path     = None
        self.document     = None
        self.step_count   = 0
        self.busy         = False
        self._settings    = {}

        self._build_ui()
        self._load_settings()
        self._on_platform_change()

        self.platform_var.trace_add("write", lambda *_: self._on_platform_change())
        self.root.bind_all("<Control-Shift-KeyPress-S>", lambda _: self._on_capture())

    # ── Build UI ─────────────────────────────────────────────

    def _build_ui(self):
        hdr = tk.Frame(self.root, bg=BG_CARD, pady=14)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Assignment Project Generator",
                 bg=BG_CARD, fg=ACCENT,
                 font=("Arial", 14, "bold")).pack()
        tk.Label(hdr, text="Screenshot  ->  AI Description  ->  Word File",
                 bg=BG_CARD, fg=TXT_SUB, font=("Arial", 9)).pack(pady=(2, 0))

        body = tk.Frame(self.root, bg=BG_DARK, padx=24, pady=14)
        body.pack(fill="both", expand=True)

        # Platform selector
        small_label(body, "LLM Platform").pack(fill="x", anchor="w")
        plat_combo = ttk.Combobox(
            body, textvariable=self.platform_var,
            values=list(PLATFORMS.keys()),
            state="readonly", font=("Arial", 10),
        )
        plat_combo.pack(fill="x", pady=(3, 12))

        # API Key
        self.key_label_var = tk.StringVar(value="API Key")
        tk.Label(body, textvariable=self.key_label_var,
                 bg=BG_DARK, fg=TXT_SUB, font=("Arial", 9), anchor="w").pack(fill="x", anchor="w")

        key_row = tk.Frame(body, bg=BG_DARK)
        key_row.pack(fill="x", pady=(3, 2))
        self.api_entry = text_entry(key_row, self.api_key_var, show="*")
        self.api_entry.pack(side="left", fill="x", expand=True)
        tk.Button(key_row, text="show", bg=ACCENT2, fg=TXT_MAIN,
                  relief="flat", cursor="hand2", font=("Arial", 8),
                  command=self._toggle_key).pack(side="left", padx=(5, 0))

        self.key_hint_lbl = tk.Label(body, text="", bg=BG_DARK, fg="#555577",
                                     font=("Arial", 8), anchor="w")
        self.key_hint_lbl.pack(fill="x", anchor="w", pady=(0, 10))

        # Model selector
        small_label(body, "Model").pack(fill="x", anchor="w")
        self.model_combo = ttk.Combobox(
            body, textvariable=self.model_var,
            state="readonly", font=("Arial", 10),
        )
        self.model_combo.pack(fill="x", pady=(3, 12))

        # Project name
        small_label(body, "Project / Assignment Name").pack(fill="x", anchor="w")
        text_entry(body, self.project_var).pack(fill="x", pady=(3, 12))

        # Subject
        small_label(body, "Subject / Context  (optional)").pack(fill="x", anchor="w")
        text_entry(body, self.subject_var).pack(fill="x", pady=(3, 12))

        # Start button
        self.start_btn = styled_button(body, "Start New Project",
                                       self._start_project, bg=ACCENT)
        self.start_btn.pack(fill="x", pady=(4, 10))

        # Divider
        tk.Frame(body, bg=ACCENT2, height=1).pack(fill="x", pady=4)

        # Step counter
        ctr = tk.Frame(body, bg=BG_DARK)
        ctr.pack(fill="x", pady=(8, 4))
        tk.Label(ctr, text="Steps captured:",
                 bg=BG_DARK, fg=TXT_SUB, font=("Arial", 10)).pack(side="left")
        self.step_lbl = tk.Label(ctr, text="0",
                                 bg=BG_DARK, fg=ACCENT,
                                 font=("Arial", 18, "bold"))
        self.step_lbl.pack(side="left", padx=10)

        # Capture button
        self.capture_btn = styled_button(
            body, "Capture Step          [Ctrl + Shift + S]",
            self._on_capture,
        )
        self.capture_btn.pack(fill="x", pady=(4, 4))
        self.capture_btn.config(state="disabled", bg="#333355")

        # Open doc button
        self.open_btn = styled_button(body, "Open Word File", self._open_doc)
        self.open_btn.pack(fill="x", pady=4)
        self.open_btn.config(state="disabled")

        # Status bar
        self.status_widget = tk.Label(
            self.root, textvariable=self.status_var,
            bg=BG_CARD, fg=TXT_SUB,
            font=("Arial", 9), wraplength=440, justify="center",
            pady=10, padx=10,
        )
        self.status_widget.pack(fill="x", side="bottom")

    # ── Platform change ──────────────────────────────────────

    def _on_platform_change(self):
        p = self.platform_var.get()
        cfg = PLATFORMS[p]
        self.key_label_var.set(cfg["key_label"])
        self.key_hint_lbl.config(text=cfg["key_hint"])
        models = cfg["models"]
        self.model_combo.config(values=models)
        saved_model = self._settings.get("models", {}).get(p, "")
        self.model_var.set(saved_model if saved_model in models else models[0])
        saved_key = self._settings.get("keys", {}).get(p, "")
        self.api_key_var.set(saved_key)

    # ── Settings ─────────────────────────────────────────────

    def _load_settings(self):
        if os.path.exists(SETTINGS_FILE):
            try:
                with open(SETTINGS_FILE) as f:
                    self._settings = json.load(f)
            except Exception:
                pass
        # Env var fallbacks
        env_map = {
            "Anthropic  (Claude)": "ANTHROPIC_API_KEY",
            "OpenAI  (GPT-4o)":    "OPENAI_API_KEY",
            "Google  (Gemini)":    "GOOGLE_API_KEY",
        }
        for plat, env_var in env_map.items():
            val = os.environ.get(env_var, "")
            if val and not self._settings.get("keys", {}).get(plat):
                self._settings.setdefault("keys", {})[plat] = val

    def _save_settings(self):
        p = self.platform_var.get()
        self._settings.setdefault("keys", {})[p]    = self.api_key_var.get()
        self._settings.setdefault("models", {})[p]  = self.model_var.get()
        try:
            with open(SETTINGS_FILE, "w") as f:
                json.dump(self._settings, f)
        except Exception:
            pass

    def _toggle_key(self):
        self.api_entry.config(
            show="" if self.api_entry.cget("show") == "*" else "*"
        )

    # ── Project start ────────────────────────────────────────

    def _start_project(self):
        key  = self.api_key_var.get().strip()
        name = self.project_var.get().strip()
        if not key:
            messagebox.showerror("Missing", "Please enter an API key for the selected platform.")
            return
        if not name:
            messagebox.showerror("Missing", "Please enter a project / assignment name.")
            return
        if not self.model_var.get():
            messagebox.showerror("Missing", "Please select a model.")
            return

        self._save_settings()

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"{name.replace(' ', '_')}_Steps.docx",
            title="Save Project File As",
        )
        if not path:
            return

        subject = self.subject_var.get().strip()
        self.document   = create_new_document(name, subject)
        self.doc_path   = path
        self.step_count = 0
        self.document.save(path)

        self.capture_btn.config(state="normal", bg=ACCENT)
        self.open_btn.config(state="normal")
        self.step_lbl.config(text="0")
        self.start_btn.config(text="Switch Project", bg=ACCENT2)
        self._set_status("Project ready. Click Capture Step or press Ctrl+Shift+S.", TXT_OK)

    # ── Capture flow ─────────────────────────────────────────

    def _on_capture(self):
        if self.busy:
            return
        if not self.document:
            messagebox.showwarning("No project", "Start a project first.")
            return
        self.busy = True
        self._set_status("Select screen region...", TXT_SUB)
        self.root.iconify()
        self.root.after(350, lambda: RegionSelector(self.root, self._on_region_selected))

    def _on_region_selected(self, image):
        self.root.deiconify()
        if image is None:
            self.busy = False
            self._set_status("Capture cancelled.", TXT_SUB)
            return
        plat  = self.platform_var.get()
        model = self.model_var.get()
        self._set_status(f"Calling {plat.strip()} / {model} ...", TXT_SUB)
        self.root.update()
        self.root.after(80, lambda: self._run_llm(image))

    def _run_llm(self, image: Image.Image):
        plat  = self.platform_var.get()
        model = self.model_var.get()
        key   = self.api_key_var.get().strip()

        try:
            description = describe_screenshot(plat, key, model, image)
        except Exception as exc:
            err = str(exc)
            if any(w in err.lower() for w in ["auth", "api key", "invalid", "401", "403"]):
                messagebox.showerror("Auth Error",
                    f"API key rejected by {plat.strip()}.\nCheck the key and try again.\n\n{err}")
            else:
                messagebox.showerror("Error", f"LLM call failed:\n{err}")
            self.busy = False
            self._set_status(f"Error: {err[:70]}", ACCENT)
            return

        self.step_count += 1
        append_step(self.document, self.step_count, description, image)
        self.document.save(self.doc_path)

        self.step_lbl.config(text=str(self.step_count))
        self._set_status(f'Step {self.step_count} added: "{description}"', TXT_OK)
        self.busy = False

    # ── Utilities ────────────────────────────────────────────

    def _open_doc(self):
        if self.doc_path and os.path.exists(self.doc_path):
            if sys.platform == "win32":
                os.startfile(self.doc_path)
            elif sys.platform == "darwin":
                os.system(f'open "{self.doc_path}"')
            else:
                os.system(f'xdg-open "{self.doc_path}"')

    def _set_status(self, msg, color=TXT_SUB):
        self.status_var.set(msg)
        self.status_widget.config(fg=color)

    def _quit(self):
        self._save_settings()
        self.root.destroy()

    def run(self):
        self.root.mainloop()


# ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    App().run()
